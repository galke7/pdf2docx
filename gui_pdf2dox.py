import os
import cv2
import docx
import pytesseract
import tkinter as tk
from tkinter import filedialog, messagebox
from PIL import Image
from pdf2image import convert_from_path
from langchain_community.document_loaders import PyPDFLoader
from langchain_core.documents.base import Document
import numpy as np

def ocr_pdf_with_preprocessing(pdf_path):
    images = convert_from_path(pdf_path)
    text = ""
    for image in images:
        open_cv_image = cv2.cvtColor(np.array(image), cv2.COLOR_RGB2BGR)
        gray = cv2.cvtColor(open_cv_image, cv2.COLOR_BGR2GRAY)
        _, thresh = cv2.threshold(gray, 128, 255, cv2.THRESH_BINARY | cv2.THRESH_OTSU)
        kernel = cv2.getStructuringElement(cv2.MORPH_RECT, (1, 1))
        opening = cv2.morphologyEx(thresh, cv2.MORPH_OPEN, kernel, iterations=1)
        text += pytesseract.image_to_string(opening)
    return text

def select_pdf_file():
    file_path = filedialog.askopenfilename(filetypes=[("PDF files", "*.pdf")])
    if file_path:
        pdf_path_entry.delete(0, tk.END)
        pdf_path_entry.insert(0, file_path)
        default_output_name = os.path.splitext(os.path.basename(file_path))[0] + ".docx"
        output_filename_entry.delete(0, tk.END)
        output_filename_entry.insert(0, default_output_name)

def select_output_location():
    file_path = filedialog.asksaveasfilename(defaultextension=".docx",
                                             filetypes=[("DOCX files", "*.docx")])
    if file_path:
        output_filename_entry.delete(0, tk.END)
        output_filename_entry.insert(0, file_path)

def convert_pdf_to_docx():
    pdf_path = pdf_path_entry.get()
    if not pdf_path:
        messagebox.showerror("Error", "Please select a PDF file.")
        return

    output_filename = output_filename_entry.get()
    if not output_filename.endswith(".docx"):
        output_filename += ".docx"

    # Check if the output filename is a full path or just a name
    if not os.path.isabs(output_filename):
        # Get the directory of the PDF file
        pdf_directory = os.path.dirname(pdf_path)
        output_filename = os.path.join(pdf_directory, output_filename)

    docs = []

    try:
        loader = PyPDFLoader(pdf_path)
        loaded_docs = loader.load()
        docs.extend(loaded_docs)
    except ValueError:
        text = ocr_pdf_with_preprocessing(pdf_path)
        doc = docx.Document()
        doc.add_paragraph(text)
        docs.append(doc)

    final_doc = docx.Document()

    for i, pdf_doc in enumerate(docs):
        if isinstance(pdf_doc, Document):
            content = pdf_doc.page_content.strip()
            if content:
                final_doc.add_paragraph(content)
        elif isinstance(pdf_doc, docx.Document):
            for para in pdf_doc.paragraphs:
                final_doc.add_paragraph(para.text)

    final_doc.save(output_filename)
    full_path = os.path.abspath(output_filename)
    messagebox.showinfo("Success", f"DOCX file created successfully as {full_path}")

# Set up the GUI
root = tk.Tk()
root.title("PDF to DOCX Converter")

# Step 1: Select PDF File
tk.Label(root, text="STEP 1 – Select PDF File 📂").grid(row=0, column=0, padx=10, pady=5, sticky="w")
pdf_path_entry = tk.Entry(root, width=50)
pdf_path_entry.grid(row=0, column=1, padx=10, pady=5)
pdf_select_button = tk.Button(root, text="Open file", command=select_pdf_file)
pdf_select_button.grid(row=0, column=2, padx=10, pady=5)

# Step 2: Select where to save new file
tk.Label(root, text="STEP 2 – Select where to save new file").grid(row=1, column=0, padx=10, pady=5, sticky="w")
output_filename_entry = tk.Entry(root, width=50)
output_filename_entry.grid(row=1, column=1, padx=10, pady=5)

# Move the "Save as..." button below the output filename entry
output_location_button = tk.Button(root, text="Save as...", command=select_output_location)
output_location_button.grid(row=1, column=2, padx=10, pady=5)

# Start converting button
convert_button = tk.Button(root, text="Start converting", command=convert_pdf_to_docx)
convert_button.grid(row=3, column=0, columnspan=3, pady=20)

root.mainloop()
